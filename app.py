import streamlit as st
from openai import OpenAI
from io import BytesIO
from docx import Document
import pandas as pd
from PyPDF2 import PdfReader

# Configuração inicial
st.set_page_config(page_title="Assisente de IA para Professores", layout="wide")

# Obtém a chave da API OpenAI dos secrets
api_key = st.secrets["OPENAI_API_KEY"]

# Configuração do cliente OpenAI
client = OpenAI(api_key=api_key)

# Listas globais para reutilização
ANOS_SERIES = [
    "Selecione uma opção",
    "EF - 1º Ano", "EF - 2º Ano", "EF - 3º Ano", "EF - 4º Ano", "EF - 5º Ano",
    "EF - 6º Ano", "EF - 7º Ano", "EF - 8º Ano", "EF - 9º Ano",
    "EM - 1º Ano", "EM - 2º Ano", "EM - 3º Ano"
]

COMPONENTES_CURRICULARES = [
    "Selecione uma opção",
    "Matemática", "Português", "Ciências", "História", "Geografia", "Arte",
    "Educação Física", "Inglês", "Biologia", "Física", "Química", "Sociologia",
    "Filosofia", "Redação", "Literatura"
]

# Função para processar uploads

def processar_arquivos(uploaded_file):
    """Processa arquivos e retorna texto extraído."""
    try:
        if uploaded_file.name.endswith(".docx"):
            doc = Document(uploaded_file)
            texto = "\n".join([p.text for p in doc.paragraphs])

        elif uploaded_file.name.endswith(".txt"):
            texto = uploaded_file.read().decode("utf-8")

        elif uploaded_file.name.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            texto = "\n".join([page.extract_text() for page in reader.pages])

        elif uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
            texto = df.to_string()

        elif uploaded_file.name.endswith(".xlsx"):
            df = pd.read_excel(uploaded_file)
            texto = df.to_string()

        else:
            st.error("Tipo de arquivo não suportado. Envie .docx, .txt, .pdf, .csv ou .xlsx.")
            return None

        return texto

    except Exception as e:
        st.error(f"Erro ao processar arquivo: {e}")
        return None

def gerar_questoes(ano, componente, assunto, dificuldade, numero_questoes, tipo, contexto=None):
    """Função para gerar questões usando a OpenAI"""
    tipo_texto = "dissertativas" if tipo == "Dissertativas" else "objetivas"
    contexto_texto = f"Utilize o seguinte contexto: \n{contexto}\n\n" if contexto else ""

    prompt = f"""
    {contexto_texto}
    Crie um conjunto de {numero_questoes} questões {tipo_texto} sobre o seguinte assunto:
    - Ano/Série: {ano}
    - Componente Curricular: {componente}
    - Assunto: {assunto if assunto else "N/A"}
    - Dificuldade: {dificuldade}

    Certifique-se de que as questões sejam claras e adequadas ao nível de ensino informado.
    """
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": "Você é um assistente especializado na criação de questões educacionais."},
                  {"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

def gerar_plano_aula(ano, componente, capitulo, modulo, duracao, metodologia, caracteristicas, assunto=None, contexto=None):
    """Função para gerar plano de aula usando a OpenAI"""
    contexto_texto = f"Utilize o seguinte contexto: \n{contexto}\n\n" if contexto else ""

    prompt = f"""
    {contexto_texto}
    Crie um plano de aula com as seguintes características:
    - Ano/Série: {ano}
    - Componente Curricular: {componente}
    - Capítulo do livro: {capitulo}
    - Módulo do capítulo: {modulo}
    - Assunto: {assunto}
    - Duração: {duracao} minutos
    - Metodologia: {metodologia}
    - Características da Turma: {caracteristicas if caracteristicas else "N/A"}
    """
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": "Você é um assistente especializado em geração de planejamento educacional para os professores."},
                  {"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content


def gerar_assunto_contextualizado(ano, componente, assunto, interesse, contexto=None):
    """Função para gerar um assunto contextualizado usando a OpenAI"""

    # Construir o início do prompt com ou sem contexto
    contexto_texto = f"Utilize o seguinte contexto: \n{contexto}\n\n" if contexto else ""

    # Construir o restante do prompt
    prompt = f"""
    {contexto_texto}
    Crie um conteúdo contextualizado com as seguintes informações:
    - Ano/Série: {ano}
    - Componente Curricular: {componente}
    - Assunto: {assunto if assunto else "N/A"}
    - Tema de Interesse: {interesse if interesse else "N/A"}
    """

    # Chamada à API
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": "Você é um assistente especializado em gerar contextualização educacional."},
                  {"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content

def corrigir_questoes(respostas_aluno, gabarito, tipo, contexto=None):
    """Função para corrigir questões usando a OpenAI"""
    tipo_texto = "dissertativas" if tipo == "Dissertativas" else "objetivas"
    contexto_texto = f"Utilize o seguinte contexto: \n{contexto}\n\n" if contexto else ""

    prompt = f"""
    {contexto_texto}
    Corrija as seguintes questões {tipo_texto} respondidas por um aluno. Baseie-se no gabarito fornecido e forneça uma análise detalhada de cada resposta:

    Respostas do Aluno:
    {respostas_aluno}

    Gabarito:
    {gabarito}

    Para cada questão, avalie:
    1. Se a resposta está correta ou não.
    2. Para questões incorretas, explique o erro e forneça a resposta correta.
    3. Para questões dissertativas, avalie a qualidade da resposta e sugira melhorias.
    """
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[{"role": "system", "content": "Você é um assistente especializado em correção de questões educacionais."},
                  {"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content


def gerar_docx(conteudo, titulo):
    """Função para gerar um arquivo DOCX a partir do conteúdo fornecido"""
    doc = Document()
    doc.add_heading(titulo, level=1)
    doc.add_paragraph(conteudo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Barra lateral
st.sidebar.title("Assistente de IA para Professores")
st.sidebar.markdown("Escolha um módulo:")
modulo = st.sidebar.selectbox("Módulos", ["Plano de Aula", "Assunto Contextualizado", "Questões", "Correção de Questões"])

# Tela principal
st.title("Assistente de IA para Professores")
st.markdown("Automatize tarefas e otimize seu planejamento educacional.")

# Verificar e processar o arquivo carregado
uploaded_file = st.sidebar.file_uploader(
    "Envie um arquivo para servir de contexto",
    type=["docx", "txt", "pdf", "csv", "xlsx"]
)

if uploaded_file:
    contexto_texto = processar_arquivos(uploaded_file)
    if contexto_texto:
        st.session_state["uploaded_file_content"] = contexto_texto
        st.success(f"Arquivo processado e armazenado para o módulo {modulo}.")

if modulo == "Plano de Aula":
    st.header("Plano de Aula")

    with st.form("plano_aula_form"):
        col1, col2 = st.columns(2)

        with col1:
            ano = st.selectbox("Ano / Série", ANOS_SERIES)
            componente = st.selectbox("Componente Curricular", COMPONENTES_CURRICULARES)
            assunto = st.text_input("Assunto", placeholder="Exemplo: Introdução à língua portuguesa")

        with col2:
            duracao = st.number_input("Duração da aula (min)", min_value=10, max_value=180, value=50)
            metodologia = st.selectbox("Metodologia", ["Selecione uma opção", "Expositiva", "Interativa", "Dinâmica"])
            caracteristicas = st.text_area("Características da Turma (opcional)", placeholder="Exemplo: Turma distraída, gosta de conversar durante a aula.")

        # Contexto do documento
        contexto = st.session_state.get("uploaded_file_content", None)

        # Mensagem indicando que "Capítulo" e "Módulo" estão desativados
        #st.markdown("⚠️ **A funcionalidade de Capítulo e Módulo estará disponível em breve.**")

        gerar = st.form_submit_button("Gerar Plano de Aula")

    if gerar:
        if ano == "Selecione uma opção" or componente == "Selecione uma opção" or metodologia == "Selecione uma opção":
            st.error("Por favor, preencha todos os campos obrigatórios!")
        else:
            with st.spinner("Gerando plano de aula..."):
                try:
                    plano_aula = gerar_plano_aula(
                        ano=ano,
                        componente=componente,
                        capitulo=None,  # Capítulo desativado temporariamente
                        modulo=None,  # Módulo desativado temporariamente
                        duracao=duracao,
                        metodologia=metodologia,
                        caracteristicas=caracteristicas,
                        assunto=assunto,
                        contexto=contexto
                    )
                    st.success("Plano de aula gerado com sucesso!")
                    st.markdown(plano_aula)

                    # Botão para download
                    buffer = gerar_docx(plano_aula, "Plano de Aula")
                    st.download_button(
                        "Baixar Plano de Aula",
                        data=buffer,
                        file_name="plano_de_aula.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Erro ao gerar plano de aula: {e}")


elif modulo == "Assunto Contextualizado":
    st.header("Assunto Contextualizado")

    # Formulário de entrada
    with st.form("contexto_form"):
        col1, col2 = st.columns(2)

        with col1:
            # Lista de anos do Ensino Fundamental e Médio
            ano = st.selectbox("Ano / Série", ANOS_SERIES)

            # Lista de componentes curriculares
            componente = st.selectbox("Componente Curricular", COMPONENTES_CURRICULARES)

            assunto = st.text_input("Assunto (opcional)", placeholder="Exemplo: Aceleração")

        with col2:
            interesse = st.text_input("Tema de Interesse (opcional)", placeholder="Exemplo: Fórmula 1")

        # Recuperar o conteúdo do arquivo carregado
        contexto = st.session_state.get("uploaded_file_content", None)

        gerar = st.form_submit_button("Gerar Assunto Contextualizado")

    if gerar:
        if ano == "Selecione uma opção" or componente == "Selecione uma opção":
            st.error("Por favor, preencha todos os campos obrigatórios!")
        else:
            with st.spinner("Gerando assunto contextualizado..."):
                try:
                    # Gerar assunto contextualizado com ou sem contexto
                    conteudo = gerar_assunto_contextualizado(ano, componente, assunto, interesse, contexto)
                    st.success("Assunto contextualizado gerado com sucesso!")
                    st.markdown(conteudo)

                    # Botão para download
                    buffer = gerar_docx(conteudo, "Assunto Contextualizado")
                    st.download_button(
                        "Baixar Assunto Contextualizado",
                        data=buffer,
                        file_name="assunto_contextualizado.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Erro ao gerar assunto contextualizado: {e}")

elif modulo == "Questões":
    st.header("Questões")

    with st.form("questoes_form"):
        col1, col2 = st.columns(2)

        with col1:
            ano = st.selectbox("Ano / Série", ANOS_SERIES)
            componente = st.selectbox("Componente Curricular", COMPONENTES_CURRICULARES)
            assunto = st.text_input("Assunto (opcional)", placeholder="Exemplo: Frações")
            numero_questoes = st.number_input("Número de Questões", min_value=1, max_value=20, value=5)

        with col2:
            dificuldade = st.selectbox("Dificuldade", ["Selecione uma opção", "Fácil", "Médio", "Difícil"])
            tipo = st.selectbox("Tipo de Questões", ["Selecione uma opção", "Objetivas", "Dissertativas"])

        contexto = st.session_state.get("uploaded_file_content", None)
        gerar = st.form_submit_button("Gerar Questões")

    if gerar:
        if ano == "Selecione uma opção" or componente == "Selecione uma opção" or dificuldade == "Selecione uma opção" or tipo == "Selecione uma opção":
            st.error("Por favor, preencha todos os campos obrigatórios!")
        else:
            with st.spinner("Gerando questões..."):
                try:
                    questoes = gerar_questoes(
                        ano, componente, assunto, dificuldade, numero_questoes, tipo, contexto
                    )
                    st.success("Questões geradas com sucesso!")
                    st.markdown(questoes)

                    # Botão para download
                    buffer = gerar_docx(questoes, "Questões")
                    st.download_button(
                        "Baixar Questões",
                        data=buffer,
                        file_name="questoes.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                except Exception as e:
                    st.error(f"Erro ao gerar questões: {e}")


elif modulo == "Correção de Questões":
    st.header("Correção de Questões")
    st.info("🚧 Este módulo estará disponível em breve! Fique ligado.")
