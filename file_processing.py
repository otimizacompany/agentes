from io import BytesIO
from docx import Document
import pandas as pd
from PyPDF2 import PdfReader
import streamlit as st

def processar_arquivos(uploaded_file):
    """Processa arquivos e retorna texto extraído."""
    try:
        if uploaded_file.name.endswith(".docx"):
            doc = Document(uploaded_file)
            texto = "\n".join([p.text for p in doc.paragraphs])

        elif uploaded_file.name.endswith(".txt"):
            texto = uploaded_file.read().decode("utf-8")

        elif uploaded_file.name.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            texto = "\n".join([page.extract_text() for page in reader.pages])

        elif uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
            texto = df.to_string()

        elif uploaded_file.name.endswith(".xlsx"):
            df = pd.read_excel(uploaded_file)
            texto = df.to_string()

        else:
            st.error("Tipo de arquivo não suportado. Envie .docx, .txt, .pdf, .csv ou .xlsx.")
            return None

        return texto

    except Exception as e:
        st.error(f"Erro ao processar arquivo: {e}")
        return None

def gerar_docx(conteudo, titulo):
    """Função para gerar um arquivo DOCX a partir do texto fornecido."""
    doc = Document()
    doc.add_heading(titulo, level=1)
    doc.add_paragraph(conteudo)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
